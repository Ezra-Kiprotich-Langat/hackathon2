import io
import logging
from fastapi import UploadFile, HTTPException
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader
from docx import Document
import chardet
from app.core.utils import clean_extracted_text, format_error_response
from app.core.config import settings

logger = logging.getLogger(__name__)

class TextExtractionService:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt,
            'jpg': self._extract_from_image,
            'jpeg': self._extract_from_image,
            'png': self._extract_from_image
        }
    
    async def extract_text(self, file: UploadFile) -> dict[str, any]:
        """Extract text from uploaded file based on its format"""
        try:
            # Get file extension
            file_extension = file.filename.split('.')[-1].lower() if file.filename else ''
            
            if file_extension not in self.supported_formats:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format: {file_extension}"
                )
            
            # Read file content
            content = await file.read()
            await file.seek(0)  # Reset file pointer
            
            # Extract text using appropriate method
            extraction_method = self.supported_formats[file_extension]
            extracted_text = await extraction_method(content)
            
            # Clean and validate extracted text
            cleaned_text = clean_extracted_text(extracted_text)
            
            if not cleaned_text.strip():
                return {
                    "success": False,
                    "text": "",
                    "word_count": 0,
                    "char_count": 0,
                    "message": "No text could be extracted from the file"
                }
            
            # Calculate statistics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "word_count": word_count,
                "char_count": char_count,
                "file_type": file_extension,
                "message": "Text extracted successfully"
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Text extraction failed for {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text: {str(e)}"
            )
    
    async def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8')
            
            # Fallback encodings if detection fails
            encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1252']
            
            for enc in encodings_to_try:
                try:
                    if enc:
                        text = content.decode(enc)
                        return text
                except (UnicodeDecodeError, LookupError):
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return content.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    async def _extract_from_image(self, content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            return text
            
        except Exception as e:
            logger.error(f"Image OCR extraction error: {str(e)}")
            # If OCR fails, return a helpful message instead of raising an error
            return "[OCR extraction failed - please ensure Tesseract is installed and configured]"
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    async def validate_text_content(self, text: str) -> dict[str, any]:
        """Validate extracted text content for question generation"""
        if not text or not text.strip():
            return {
                "valid": False,
                "reason": "No text content found"
            }
        
        word_count = len(text.split())
        
        if word_count < 50:
            return {
                "valid": False,
                "reason": "Text too short for meaningful question generation (minimum 50 words)"
            }
        
        if word_count > 10000:
            return {
                "valid": False,
                "reason": "Text too long for processing (maximum 10,000 words)"
            }
        
        return {
            "valid": True,
            "word_count": word_count,
            "char_count": len(text)
        }

# Create service instance
text_extraction_service = TextExtractionService()